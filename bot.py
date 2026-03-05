"""
📄 Telegram → AI → DOCX Bot
Completely FREE stack:
  - Telegram Bot API (free)
  - Google Gemini 1.5 Flash (free tier: 1500 req/day)
  - python-docx (free library)
"""

import os
import io
import logging
import tempfile
import base64
import re
from pathlib import Path

import google.generativeai as genai
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from telegram import Update
from telegram.ext import (
    Application,
    CommandHandler,
    MessageHandler,
    filters,
    ContextTypes,
)

# ─── Logging ──────────────────────────────────────────────────────────────────
logging.basicConfig(
    format="%(asctime)s | %(levelname)s | %(message)s", level=logging.INFO
)
logger = logging.getLogger(__name__)

# ─── Config (set via environment variables) ───────────────────────────────────
TELEGRAM_TOKEN = os.environ["TELEGRAM_TOKEN"]   # From @BotFather
GEMINI_API_KEY  = os.environ["GEMINI_API_KEY"]  # From aistudio.google.com (free)

genai.configure(api_key=GEMINI_API_KEY)
model = genai.GenerativeModel("gemini-1.5-flash")

EXTRACTION_PROMPT = """You are a precise document extraction assistant.

Extract ALL content from this file with EXACT formatting preserved:
- Keep all headings, subheadings, and section titles
- Preserve bullet points and numbered lists exactly
- Maintain tables with their structure (use | to separate columns)
- Keep all bold, italic indicators if visible
- Preserve paragraph spacing and indentation structure
- Extract every word — do NOT summarize or skip anything

Return the content in clean plain text with formatting markers:
  # for main headings
  ## for subheadings  
  ### for sub-subheadings
  - for bullet points
  1. for numbered lists
  **text** for bold
  *text* for italic
  TABLE: for table rows using | separator
  
Return ONLY the extracted content, nothing else."""


# ─── Helpers ──────────────────────────────────────────────────────────────────

async def download_file_bytes(file) -> bytes:
    """Download a Telegram file object and return raw bytes."""
    file_obj = await file.get_file()
    buf = io.BytesIO()
    await file_obj.download_to_memory(buf)
    return buf.getvalue()


def extract_with_gemini(file_bytes: bytes, mime_type: str) -> str:
    """Send file to Gemini and get extracted text back."""
    part = {
        "inline_data": {
            "mime_type": mime_type,
            "data": base64.b64encode(file_bytes).decode(),
        }
    }
    response = model.generate_content([EXTRACTION_PROMPT, part])
    return response.text


def build_docx(extracted_text: str) -> bytes:
    """Convert extracted markdown-style text into a formatted DOCX."""
    doc = Document()

    # ── Document styles ──────────────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    # ── Process each line ────────────────────────────────────────────────────
    lines = extracted_text.split("\n")
    i = 0

    while i < len(lines):
        raw = lines[i]
        line = raw.strip()

        # Skip blanks (add spacing via paragraph spacing instead)
        if not line:
            i += 1
            continue

        # Headings
        if line.startswith("### "):
            p = doc.add_heading(line[4:], level=3)
        elif line.startswith("## "):
            p = doc.add_heading(line[3:], level=2)
        elif line.startswith("# "):
            p = doc.add_heading(line[2:], level=1)

        # Tables  (TABLE: col1 | col2 | col3)
        elif line.upper().startswith("TABLE:"):
            cols = [c.strip() for c in line[6:].split("|")]
            # Peek ahead to collect all table rows
            table_rows = [cols]
            j = i + 1
            while j < len(lines):
                next_line = lines[j].strip()
                if "|" in next_line and not next_line.startswith("#"):
                    table_rows.append([c.strip() for c in next_line.split("|")])
                    j += 1
                else:
                    break
            i = j - 1  # will be incremented at end of loop

            max_cols = max(len(r) for r in table_rows)
            tbl = doc.add_table(rows=len(table_rows), cols=max_cols)
            tbl.style = "Table Grid"
            for r_idx, row_data in enumerate(table_rows):
                for c_idx, cell_text in enumerate(row_data):
                    cell = tbl.rows[r_idx].cells[c_idx]
                    cell.text = cell_text
                    if r_idx == 0:  # Bold header row
                        for run in cell.paragraphs[0].runs:
                            run.bold = True

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, line[2:])

        # Numbered list
        elif re.match(r"^\d+\.\s", line):
            p = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, re.sub(r"^\d+\.\s", "", line))

        # Normal paragraph
        else:
            p = doc.add_paragraph()
            _add_inline_formatting(p, line)

        i += 1

    # ── Save to bytes ─────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


def _add_inline_formatting(paragraph, text: str):
    """Parse **bold** and *italic* markers and add styled runs."""
    # Split on bold/italic markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


# ─── Bot handlers ──────────────────────────────────────────────────────────────

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "👋 Hello! I'm your *Document Extractor Bot*.\n\n"
        "📤 Send me any *image* (JPG, PNG) or *PDF* and I'll:\n"
        "  1️⃣ Extract all the content using AI\n"
        "  2️⃣ Generate a formatted *.docx* file\n"
        "  3️⃣ Send it right back to you!\n\n"
        "Just drop your file here to get started 🚀",
        parse_mode="Markdown",
    )


async def handle_document(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle PDF or image files sent as documents."""
    doc = update.message.document
    mime = doc.mime_type or ""

    if "pdf" in mime:
        await _process_file(update, context, doc, "application/pdf", "document.pdf")
    elif "image" in mime:
        await _process_file(update, context, doc, mime, doc.file_name or "image.jpg")
    else:
        await update.message.reply_text(
            "⚠️ Please send a *PDF* or *image* file (JPG, PNG, etc.).",
            parse_mode="Markdown",
        )


async def handle_photo(update: Update, context: ContextTypes.DEFAULT_TYPE):
    """Handle photos sent directly (compressed by Telegram)."""
    photo = update.message.photo[-1]  # highest resolution
    await _process_file(update, context, photo, "image/jpeg", "photo.jpg")


async def _process_file(update, context, file_obj, mime_type: str, filename: str):
    """Core pipeline: download → AI extract → build DOCX → send back."""
    status_msg = await update.message.reply_text("⏳ Downloading your file…")

    try:
        # 1. Download
        file_bytes = await download_file_bytes(file_obj)
        await status_msg.edit_text("🤖 Extracting content with AI…")

        # 2. AI extraction
        extracted_text = extract_with_gemini(file_bytes, mime_type)
        logger.info("Extracted %d chars from %s", len(extracted_text), filename)
        await status_msg.edit_text("📝 Building your DOCX file…")

        # 3. Build DOCX
        docx_bytes = build_docx(extracted_text)
        output_name = Path(filename).stem + "_extracted.docx"

        # 4. Send DOCX back
        await status_msg.edit_text("📤 Sending your document…")
        await update.message.reply_document(
            document=io.BytesIO(docx_bytes),
            filename=output_name,
            caption=f"✅ Done! Extracted from *{filename}*",
            parse_mode="Markdown",
        )
        await status_msg.delete()

    except Exception as e:
        logger.exception("Processing failed")
        await status_msg.edit_text(
            f"❌ Something went wrong:\n`{str(e)}`\n\nPlease try again.",
            parse_mode="Markdown",
        )


async def handle_unsupported(update: Update, context: ContextTypes.DEFAULT_TYPE):
    await update.message.reply_text(
        "Please send a *PDF* or *image* file.\nType /start for instructions.",
        parse_mode="Markdown",
    )


# ─── Main ──────────────────────────────────────────────────────────────────────

def main():
    app = Application.builder().token(TELEGRAM_TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_document))
    app.add_handler(MessageHandler(filters.PHOTO, handle_photo))
    app.add_handler(MessageHandler(filters.ALL & ~filters.COMMAND, handle_unsupported))

    logger.info("🤖 Bot is running…")
    app.run_polling()


if __name__ == "__main__":
    main()
