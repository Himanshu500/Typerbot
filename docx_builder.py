"""
docx_builder.py — Full DOCX builder with:
  - Sentence merge engine
  - Table reconstruction
  - Chemical structure renderer
  - Superscript/subscript engine
  - Two-column side-by-side renderer
  - Signature table renderer
  - Full style application (font, size, spacing, alignment)
"""

import io
import re
import logging
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from style_engine import StyleMap, ElementStyle

logger = logging.getLogger(__name__)

# ── Unicode sub/superscript maps ──────────────────────────────────────────────
SUB_MAP  = str.maketrans("0123456789", "₀₁₂₃₄₅₆₇₈₉")
SUP_ORDS = {"th", "st", "nd", "rd"}
CHEM_SUB_RE = re.compile(r'([A-Za-z)])(\d+)')   # CH3 → CH₃


# ══════════════════════════════════════════════════════════════════════════════
#  PUBLIC ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def build_docx(extracted_text: str, style_map: StyleMap) -> bytes:
    """
    Convert AI-extracted marked-up text into a formatted DOCX.
    """
    doc = Document()
    _setup_margins(doc)

    lines = extracted_text.split("\n")
    lines = _merge_sentences(lines)          # Step 1: fix line-break splits
    blocks = _parse_blocks(lines)            # Step 2: group into semantic blocks
    _render_blocks(doc, blocks, style_map)   # Step 3: render to DOCX

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()


# ══════════════════════════════════════════════════════════════════════════════
#  STEP 1 — SENTENCE MERGE ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def _merge_sentences(lines: list[str]) -> list[str]:
    """
    Merge lines that are physical wraps, not logical paragraph breaks.
    Rules:
      - Line ends without .!?:; AND next line starts lowercase → merge
      - Short orphan line (< 5 words) after a list item → re-attach
      - Hyphen at end → word-break, merge and remove hyphen
    """
    result = []
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        # Skip blank lines — preserve as paragraph separators
        if not stripped.strip():
            result.append(line)
            i += 1
            continue

        # Look ahead to see if we should merge
        while i + 1 < len(lines):
            next_line = lines[i + 1].strip()

            # Don't merge across blank lines
            if not next_line:
                break

            # Don't merge into marker lines
            if _is_marker(next_line):
                break

            # Don't merge if current line is a marker
            if _is_marker(stripped.strip()):
                break

            # Hyphen word-break: merge and remove hyphen
            if stripped.endswith("-") and next_line and next_line[0].islower():
                stripped = stripped[:-1] + next_line
                i += 1
                continue

            # Line ends without sentence terminator AND next starts lowercase
            current_text = stripped.strip()
            ends_incomplete = current_text and current_text[-1] not in ".!?:;)]"
            next_starts_lower = next_line and (next_line[0].islower() or next_line[0] in "abcdefghijklmnopqrstuvwxyz,")

            if ends_incomplete and next_starts_lower:
                stripped = stripped.rstrip() + " " + next_line
                i += 1
                continue

            break

        result.append(stripped)
        i += 1

    return result


def _is_marker(line: str) -> bool:
    """Check if line is a special format marker — should not be merged."""
    markers = (
        "CENTER:", "REFDATE:", "BOLD:", "BOLDSTART:", "SUBJECT:", "BOLDNUM:",
        "SIGTABLE:", "SIGNAME:", "SIGROLE:", "COMMENTBOX:", "FOOTNOTE:",
        "SIDEBY_LEFT:", "SIDEBY_RIGHT:", "QA_TABLE:", "END_QA_TABLE",
        "STRUCT:", "END_STRUCT", "TABLE:", "END_TABLE", "FIELD:",
        "#", "##", "###", "- ", "1.", "2.", "3.", "4.", "5.",
        "[LOGO:", "[unclear:", "[handwritten:",
    )
    s = line.strip()
    return any(s.startswith(m) for m in markers)


# ══════════════════════════════════════════════════════════════════════════════
#  STEP 2 — BLOCK PARSER
# ══════════════════════════════════════════════════════════════════════════════

def _parse_blocks(lines: list[str]) -> list[dict]:
    """Parse lines into semantic blocks for rendering."""
    blocks = []
    i = 0

    while i < len(lines):
        line = lines[i].strip()

        if not line:
            blocks.append({"type": "blank"})
            i += 1
            continue

        # ── Structural formula block ──────────────────────────────────────────
        if line == "STRUCT:":
            struct_lines = []
            i += 1
            while i < len(lines) and lines[i].strip() != "END_STRUCT":
                struct_lines.append(lines[i])
                i += 1
            blocks.append({"type": "struct", "lines": struct_lines})
            i += 1
            continue

        # ── Q&A table (science textbook two-column) ───────────────────────────
        if line == "QA_TABLE:":
            rows = []
            i += 1
            while i < len(lines) and lines[i].strip() != "END_QA_TABLE":
                row_line = lines[i].strip()
                if "|" in row_line:
                    parts = [p.strip() for p in row_line.split("|", 1)]
                    rows.append(parts)
                i += 1
            blocks.append({"type": "qa_table", "rows": rows})
            i += 1
            continue

        # ── Regular table ─────────────────────────────────────────────────────
        if line == "TABLE:":
            rows = []
            i += 1
            while i < len(lines) and lines[i].strip() != "END_TABLE":
                row_line = lines[i].strip()
                if row_line:
                    # Strip stray "TABLE:" prefix (known AI bug)
                    row_line = re.sub(r'^TABLE:\s*', '', row_line)
                    cols = [c.strip() for c in row_line.split("|")]
                    rows.append(cols)
                i += 1
            if rows:
                blocks.append({"type": "table", "rows": rows})
            i += 1
            continue

        # ── Side-by-side blocks ───────────────────────────────────────────────
        if line.startswith("SIDEBY_LEFT:"):
            left_lines  = [line[12:].strip()]
            right_lines = []
            i += 1
            while i < len(lines) and lines[i].strip().startswith("SIDEBY_LEFT:"):
                left_lines.append(lines[i].strip()[12:].strip())
                i += 1
            while i < len(lines) and lines[i].strip().startswith("SIDEBY_RIGHT:"):
                right_lines.append(lines[i].strip()[13:].strip())
                i += 1
            blocks.append({"type": "sideby", "left": left_lines, "right": right_lines})
            continue

        # ── Signature table ───────────────────────────────────────────────────
        if line == "SIGTABLE:":
            sig_data = []   # list of {header, name, role}
            i += 1
            current = {}
            while i < len(lines):
                sl = lines[i].strip()
                if sl.startswith("SIGNAME:"):
                    current["name"] = sl[8:].strip()
                elif sl.startswith("SIGROLE:"):
                    current["role"] = sl[8:].strip()
                elif sl.startswith("SIGHEADER:"):
                    if current:
                        sig_data.append(current)
                    current = {"header": sl[10:].strip()}
                elif sl in ("END_SIGTABLE", "") and current:
                    sig_data.append(current)
                    current = {}
                    if sl == "END_SIGTABLE":
                        break
                else:
                    break
                i += 1
            if current:
                sig_data.append(current)
            blocks.append({"type": "sigtable", "cols": sig_data})
            continue

        # ── Comment box ───────────────────────────────────────────────────────
        if line.startswith("COMMENTBOX:"):
            blocks.append({"type": "commentbox", "title": line[11:].strip()})
            i += 1
            continue

        # ── Ref/Date line ─────────────────────────────────────────────────────
        if line.startswith("REFDATE:"):
            parts = line[8:].split("|", 1)
            blocks.append({
                "type": "refdate",
                "ref":  parts[0].strip() if parts else "",
                "date": parts[1].strip() if len(parts) > 1 else "",
            })
            i += 1
            continue

        # ── Headings ──────────────────────────────────────────────────────────
        if line.startswith("### "):
            blocks.append({"type": "h3", "text": line[4:]})
        elif line.startswith("## "):
            blocks.append({"type": "h2", "text": line[3:]})
        elif line.startswith("# "):
            blocks.append({"type": "h1", "text": line[2:]})

        # ── CENTER ────────────────────────────────────────────────────────────
        elif line.startswith("CENTER:"):
            blocks.append({"type": "center", "text": line[7:].strip()})

        # ── Subject line ──────────────────────────────────────────────────────
        elif line.startswith("SUBJECT:"):
            blocks.append({"type": "subject", "text": line[8:].strip()})

        # ── Bold number (section label) ───────────────────────────────────────
        elif line.startswith("BOLDNUM:"):
            blocks.append({"type": "boldnum", "text": line[8:].strip()})

        # ── Footnote ──────────────────────────────────────────────────────────
        elif line.startswith("FOOTNOTE:"):
            blocks.append({"type": "footnote", "text": line[9:].strip()})

        # ── Field (form) ──────────────────────────────────────────────────────
        elif line.startswith("FIELD:"):
            parts = line[6:].split("|", 1)
            blocks.append({
                "type": "field",
                "label": parts[0].strip(),
                "value": parts[1].strip() if len(parts) > 1 else "",
            })

        # ── Bullet list ───────────────────────────────────────────────────────
        elif line.startswith("- ") or line.startswith("* "):
            blocks.append({"type": "bullet", "text": line[2:]})

        # ── Numbered list ─────────────────────────────────────────────────────
        elif re.match(r"^(\d{1,2}[.)]\s|[ivxIVX]+[.)]\s)", line):
            # Preserve original number format
            m = re.match(r"^(\d{1,2}[.)]\s|[ivxIVX]+[.)]\s)(.*)", line)
            if m:
                blocks.append({"type": "numbered", "num": m.group(1).strip(), "text": m.group(2)})
            else:
                blocks.append({"type": "numbered", "num": "", "text": line})

        # ── Logo marker ───────────────────────────────────────────────────────
        elif line.startswith("[LOGO:"):
            blocks.append({"type": "logo", "text": line})

        # ── Normal paragraph ──────────────────────────────────────────────────
        else:
            blocks.append({"type": "para", "text": line})

        i += 1

    # ── Post-process: merge consecutive 1-col table blocks ───────────────────
    blocks = _merge_solo_tables(blocks)

    return blocks


def _merge_solo_tables(blocks: list[dict]) -> list[dict]:
    """Merge sequences of single-row tables into one multi-row table."""
    result = []
    i = 0
    while i < len(blocks):
        b = blocks[i]
        if b["type"] == "table" and len(b["rows"]) == 1 and len(b["rows"][0]) == 1:
            # Collect consecutive single-cell tables
            merged_rows = [b["rows"][0]]
            j = i + 1
            while j < len(blocks):
                nb = blocks[j]
                if nb["type"] == "table" and len(nb["rows"]) == 1 and len(nb["rows"][0]) == 1:
                    merged_rows.append(nb["rows"][0])
                    j += 1
                else:
                    break
            if len(merged_rows) > 1:
                result.append({"type": "table", "rows": merged_rows})
                i = j
                continue
        result.append(b)
        i += 1
    return result


# ══════════════════════════════════════════════════════════════════════════════
#  STEP 3 — RENDERER
# ══════════════════════════════════════════════════════════════════════════════

def _render_blocks(doc: Document, blocks: list[dict], sm: StyleMap):
    """Render all blocks to the document."""
    for b in blocks:
        t = b["type"]

        if t == "blank":
            pass   # skip — spacing handled by paragraph spacing

        elif t == "h1":
            _render_heading(doc, b["text"], 1, sm)

        elif t == "h2":
            _render_heading(doc, b["text"], 2, sm)

        elif t == "h3":
            _render_heading(doc, b["text"], 3, sm)

        elif t == "center":
            _render_center(doc, b["text"], sm)

        elif t == "subject":
            _render_subject(doc, b["text"], sm)

        elif t == "refdate":
            _render_refdate(doc, b["ref"], b["date"], sm)

        elif t == "boldnum":
            _render_boldnum(doc, b["text"], sm)

        elif t == "para":
            _render_para(doc, b["text"], sm)

        elif t == "bullet":
            _render_list_item(doc, b["text"], sm, numbered=False)

        elif t == "numbered":
            _render_list_item(doc, b["text"], sm, numbered=True, num=b.get("num", ""))

        elif t == "table":
            _render_table(doc, b["rows"], sm)

        elif t == "qa_table":
            _render_qa_table(doc, b["rows"], sm)

        elif t == "struct":
            _render_struct(doc, b["lines"], sm)

        elif t == "sideby":
            _render_sideby(doc, b["left"], b["right"], sm)

        elif t == "sigtable":
            _render_sigtable(doc, b["cols"], sm)

        elif t == "commentbox":
            _render_commentbox(doc, b["title"], sm)

        elif t == "footnote":
            _render_footnote(doc, b["text"], sm)

        elif t == "field":
            _render_field(doc, b["label"], b["value"], sm)

        elif t == "logo":
            _render_logo_note(doc, b["text"], sm)


# ── Individual renderers ──────────────────────────────────────────────────────

def _render_heading(doc, text, level, sm):
    role_map = {1: "main_title", 2: "heading", 3: "subheading"}
    role = role_map.get(level, "heading")
    es   = sm.get(role)
    p    = doc.add_heading(level=level)
    p.clear()
    _apply_para_fmt(p, es)
    _add_inline_text(p, text, es)


def _render_center(doc, text, sm):
    # Could be title or subtitle — pick by content
    is_big = text.isupper() or len(text) < 50
    role   = "main_title" if is_big else "subtitle"
    es     = sm.get(role)
    es_copy = ElementStyle(
        font_name=es.font_name, font_size=es.font_size,
        bold=es.bold, italic=es.italic,
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_before=es.space_before, space_after=es.space_after,
        line_spacing=es.line_spacing,
    )
    p = doc.add_paragraph()
    _apply_para_fmt(p, es_copy)
    _add_inline_text(p, text, es_copy)


def _render_subject(doc, text, sm):
    es = sm.get("subject_line")
    p  = doc.add_paragraph()
    _apply_para_fmt(p, es)
    _add_inline_text(p, text, es)


def _render_refdate(doc, ref, date, sm):
    """Ref number left, date right — using tab stop."""
    from docx.oxml.ns import qn
    es = sm.get("ref_line")
    p  = doc.add_paragraph()
    _apply_para_fmt(p, es)

    # Left part
    run_ref = p.add_run(ref)
    _style_run(run_ref, es)

    # Tab to right margin
    run_tab = p.add_run("\t")
    _style_run(run_tab, es)

    # Right-aligned date
    run_date = p.add_run(date)
    _style_run(run_date, es)
    run_date.bold = True

    # Add right-align tab stop at 6 inches
    from docx.oxml import OxmlElement
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "8640")   # 6 inches in twips
    tabs.append(tab)
    pPr.append(tabs)


def _render_boldnum(doc, text, sm):
    es = sm.get("heading")
    p  = doc.add_paragraph()
    _apply_para_fmt(p, es)
    run = p.add_run(text)
    _style_run(run, es)
    run.bold = True


def _render_para(doc, text, sm):
    es = sm.get("body")
    p  = doc.add_paragraph()
    _apply_para_fmt(p, es)
    _add_inline_text(p, text, es)


def _render_list_item(doc, text, sm, numbered=False, num=""):
    es    = sm.get("list_item")
    style = "List Number" if numbered else "List Bullet"
    p     = doc.add_paragraph(style=style)
    _apply_para_fmt(p, es)
    _add_inline_text(p, text, es)


def _render_table(doc, rows, sm):
    if not rows:
        return
    max_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=max_cols)
    tbl.style = "Table Grid"

    es_header = sm.get("table_header")
    es_cell   = sm.get("table_cell")

    for r_idx, row_data in enumerate(rows):
        is_header = r_idx == 0
        for c_idx in range(max_cols):
            cell_text = row_data[c_idx] if c_idx < len(row_data) else ""
            cell = tbl.rows[r_idx].cells[c_idx]
            _fill_cell(cell, cell_text, es_header if is_header else es_cell)

    _set_table_margins(tbl)


def _render_qa_table(doc, rows, sm):
    """Two-column Q&A table for science textbook."""
    if not rows:
        return
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.style = "Table Grid"

    es_struct = sm.get("struct")
    es_body   = sm.get("body")

    # Set column widths: 40% formula, 60% explanation
    for row_data in rows:
        formula     = row_data[0] if len(row_data) > 0 else ""
        explanation = row_data[1] if len(row_data) > 1 else ""
        r_idx = rows.index(row_data)
        row   = tbl.rows[r_idx]
        _fill_cell(row.cells[0], formula, es_struct)
        _fill_cell(row.cells[1], explanation, es_body)

    _set_table_margins(tbl)

    # Set column widths
    from docx.shared import Inches
    for row in tbl.rows:
        row.cells[0].width = Inches(2.5)
        row.cells[1].width = Inches(4.0)


def _render_struct(doc, lines, sm):
    """Render chemical structural formula in monospace box."""
    es = sm.get("struct")
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = "Table Grid"
    cell = tbl.rows[0].cells[0]

    # Gray background
    _set_cell_shading(cell, "F2F2F2")

    # Remove existing paragraphs and add each line
    for i, formula_line in enumerate(lines):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        _apply_para_fmt(p, es)
        run = p.add_run(formula_line)
        _style_run(run, es)
        run.font.name = "Courier New"
        run.font.size = Pt(10)

    _set_table_margins(tbl)


def _render_sideby(doc, left_lines, right_lines, sm):
    """Side-by-side two-column borderless table."""
    es = sm.get("body")
    tbl = doc.add_table(rows=1, cols=2)
    # No borders
    for row in tbl.rows:
        for cell in row.cells:
            _remove_cell_borders(cell)

    left_cell  = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]

    for i, line in enumerate(left_lines):
        p = left_cell.paragraphs[0] if i == 0 else left_cell.add_paragraph()
        _apply_para_fmt(p, es)
        _add_inline_text(p, line, es)

    for i, line in enumerate(right_lines):
        p = right_cell.paragraphs[0] if i == 0 else right_cell.add_paragraph()
        _apply_para_fmt(p, es)
        _add_inline_text(p, line, es)

    # Column widths
    left_cell.width  = Inches(3.0)
    right_cell.width = Inches(3.5)


def _render_sigtable(doc, cols, sm):
    """3-column signature table with bold names."""
    if not cols:
        return
    num_cols = len(cols)
    tbl = doc.add_table(rows=3, cols=num_cols)   # header | name | role rows
    tbl.style = "Table Grid"

    es_header = sm.get("table_header")
    es_name   = sm.get("signature_name")
    es_role   = sm.get("signature_role")

    for c_idx, col in enumerate(cols):
        _fill_cell(tbl.rows[0].cells[c_idx], col.get("header", ""), es_header)
        _fill_cell(tbl.rows[1].cells[c_idx], col.get("name", ""), es_name)
        _fill_cell(tbl.rows[2].cells[c_idx], col.get("role", "[Signature]"), es_role)

    _set_table_margins(tbl)


def _render_commentbox(doc, title, sm):
    """Bordered box with title header and empty body for comments."""
    tbl = doc.add_table(rows=2, cols=1)
    tbl.style = "Table Grid"
    es_header = sm.get("table_header")
    es_body   = sm.get("body")
    _fill_cell(tbl.rows[0].cells[0], title, es_header)
    # Empty body row
    p = tbl.rows[1].cells[0].paragraphs[0]
    _apply_para_fmt(p, es_body)
    p.add_run("")
    # Make body row taller
    from docx.shared import Inches
    tbl.rows[1].height = Inches(0.6)
    _set_table_margins(tbl)


def _render_footnote(doc, text, sm):
    es = sm.get("footnote")
    p  = doc.add_paragraph()
    _apply_para_fmt(p, es)
    _add_inline_text(p, text, es)


def _render_field(doc, label, value, sm):
    """Form field: bold label + value on same line."""
    es_label = sm.get("label")
    es_body  = sm.get("body")
    p = doc.add_paragraph()
    _apply_para_fmt(p, es_body)
    run_label = p.add_run(label + ": ")
    _style_run(run_label, es_label)
    run_label.bold = True
    run_value = p.add_run(value)
    _style_run(run_value, es_body)


def _render_logo_note(doc, text, sm):
    es = sm.get("caption")
    p  = doc.add_paragraph()
    _apply_para_fmt(p, es)
    p.add_run(text).italic = True


# ══════════════════════════════════════════════════════════════════════════════
#  INLINE TEXT ENGINE
# ══════════════════════════════════════════════════════════════════════════════

def _add_inline_text(paragraph, text: str, base_es: ElementStyle):
    """
    Parse inline markers and add styled runs.
    Handles: **bold**, *italic*, BOLD: prefix, BOLDSTART: prefix,
             ^superscript^, ordinal superscripts, chemical subscripts.
    """
    # Pre-process: remove BOLDSTART: / BOLD: markers
    text = re.sub(r'^BOLDSTART:\s*', '', text.strip())
    text = re.sub(r'^BOLD:\s*', '', text.strip())

    # Split on inline markers: **bold**, *italic*, ^super^
    tokens = re.split(r'(\*\*[^*]+\*\*|\*[^*]+\*|\^[^^]+\^)', text)

    for token in tokens:
        if not token:
            continue

        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(_process_subscripts(token[2:-2]))
            _style_run(run, base_es)
            run.bold = True

        elif token.startswith("*") and token.endswith("*"):
            run = paragraph.add_run(_process_subscripts(token[1:-1]))
            _style_run(run, base_es)
            run.italic = True

        elif token.startswith("^") and token.endswith("^"):
            # Superscript
            run = paragraph.add_run(token[1:-1])
            _style_run(run, base_es)
            _set_superscript(run)

        else:
            # Process ordinal superscripts (26th → 26 + th superscript)
            _add_with_ordinals(paragraph, token, base_es)


def _add_with_ordinals(paragraph, text: str, base_es: ElementStyle):
    """Split text at ordinal suffixes and render th/st/nd/rd as superscript."""
    # Match: digits followed by th/st/nd/rd
    parts = re.split(r'(\d+(?:st|nd|rd|th))', text)
    for part in parts:
        if not part:
            continue
        m = re.match(r'^(\d+)(st|nd|rd|th)$', part)
        if m:
            # Number part normal
            run_num = paragraph.add_run(_process_subscripts(m.group(1)))
            _style_run(run_num, base_es)
            # Suffix as superscript
            run_sup = paragraph.add_run(m.group(2))
            _style_run(run_sup, base_es)
            _set_superscript(run_sup)
        else:
            run = paragraph.add_run(_process_subscripts(part))
            _style_run(run, base_es)


def _process_subscripts(text: str) -> str:
    """Convert chemical formula digit subscripts: CH3 → CH₃"""
    return CHEM_SUB_RE.sub(lambda m: m.group(1) + m.group(2).translate(SUB_MAP), text)


def _set_superscript(run):
    """Apply superscript via XML."""
    rPr = run._r.get_or_add_rPr()
    vertAlign = OxmlElement("w:vertAlign")
    vertAlign.set(qn("w:val"), "superscript")
    rPr.append(vertAlign)


def _set_subscript(run):
    """Apply subscript via XML."""
    rPr = run._r.get_or_add_rPr()
    vertAlign = OxmlElement("w:vertAlign")
    vertAlign.set(qn("w:val"), "subscript")
    rPr.append(vertAlign)


# ══════════════════════════════════════════════════════════════════════════════
#  STYLE APPLICATION HELPERS
# ══════════════════════════════════════════════════════════════════════════════

def _apply_para_fmt(paragraph, es: ElementStyle):
    fmt = paragraph.paragraph_format
    if es.align is not None:
        fmt.alignment     = es.align
    if es.space_before is not None:
        fmt.space_before  = es.space_before
    if es.space_after is not None:
        fmt.space_after   = es.space_after
    if es.line_spacing is not None:
        fmt.line_spacing  = es.line_spacing


def _style_run(run, es: ElementStyle):
    if es.font_name:
        run.font.name = es.font_name
    if es.font_size:
        run.font.size = Pt(es.font_size)
    if es.bold:
        run.bold = True
    if es.italic:
        run.italic = True
    if es.color:
        run.font.color.rgb = es.color


def _fill_cell(cell, text: str, es: ElementStyle):
    p = cell.paragraphs[0]
    p.clear()
    _apply_para_fmt(p, es)
    _add_inline_text(p, text, es)


def _set_cell_shading(cell, hex_color: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _remove_cell_borders(cell):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        border = OxmlElement(f"w:{side}")
        border.set(qn("w:val"),  "none")
        border.set(qn("w:sz"),   "0")
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "auto")
        tcBorders.append(border)
    tcPr.append(tcBorders)


def _set_table_margins(tbl):
    """Add cell padding to all cells."""
    tbl_elem = tbl._tbl
    # Get or create tblPr manually
    tblPr = tbl_elem.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl_elem.insert(0, tblPr)
    # Remove existing tblCellMar if present
    existing = tblPr.find(qn("w:tblCellMar"))
    if existing is not None:
        tblPr.remove(existing)
    tblCellMar = OxmlElement("w:tblCellMar")
    for side, val in [("top", "80"), ("bottom", "80"), ("left", "120"), ("right", "120")]:
        node = OxmlElement(f"w:{side}")
        node.set(qn("w:w"),    val)
        node.set(qn("w:type"), "dxa")
        tblCellMar.append(node)
    tblPr.append(tblCellMar)


def _setup_margins(doc):
    """Set standard 1-inch margins."""
    from docx.shared import Inches
    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1)
        section.right_margin  = Inches(1)
